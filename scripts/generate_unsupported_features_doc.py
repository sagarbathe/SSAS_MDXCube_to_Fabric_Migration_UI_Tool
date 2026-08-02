"""
One-off generator for docs/Unsupported_Features_Feasibility.docx - a
feasibility write-up on whether/how the currently-unsupported SSAS
Multidimensional features (Row-Level Security, Actions, Perspectives,
Translations, custom rollups/unary operators, write-back, MDX SCOPE
assignments) could be added to this migration tool. This is a review
document, not an implementation plan that's been committed to - nothing
in the pipeline changes as a result of generating it.

Not part of the pipeline; run manually whenever this content needs to be
regenerated:
    .venv-ui\\Scripts\\python.exe scripts\\generate_unsupported_features_doc.py
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(REPO_ROOT, "docs", "Unsupported_Features_Feasibility.docx")

SUMMARY_ROWS = [
    ("Feature", "Feasible to add?", "Effort", "Fabric/DAX target support"),
    ("Row-Level Security (roles)", "Yes", "Medium", "Fully supported"),
    ("Perspectives", "Yes", "Low", "Supported, but not honored by Direct Lake reports today"),
    ("Translations", "Partial", "Low-Medium", "Limited support in Direct Lake"),
    ("Actions (drillthrough/reporting)", "No (not 1:1)", "N/A", "Not a TMDL/Fabric concept"),
    ("Custom rollup / unary operators", "Partial (workaround only)", "High", "No direct DAX equivalent"),
    ("Write-back", "No", "N/A", "Fabric semantic models are read-only"),
    ("MDX SCOPE assignments", "Partial (case-by-case)", "High", "No direct DAX equivalent; needs manual redesign"),
]

# Each entry: (heading_level, "para"|"bullet", text)
SECTIONS = [
    (1, "heading", "1. Row-Level Security (Roles)"),
    (0, "para", "Feasible - recommend doing first."),
    (0, "bullet",
     "SSAS AMO exposes Database.Roles (name, members, and per-cube CubePermission "
     "with RowFilterID - an MDX filter set expression)."),
    (0, "bullet",
     "TMDL's equivalent is a role { modelPermission = \"read\" tablePermission "
     "'Table' { filterExpression = \"<DAX>\" } } block per table."),
    (0, "bullet",
     "Work needed: (a) extractor - pull the Roles collection, per-cube "
     "permissions, and the MDX row-filter expression into a new RoleIR; (b) a "
     "translation step from the MDX filter (e.g. [Dim].[Attr].&[Value]-style "
     "tuple/set filters) to an equivalent DAX boolean expression - straightforward "
     "for simple attribute-based filters, harder for nested/dynamic MDX filters; "
     "(c) TMDL generator emits a definition/roles/<RoleName>.tmdl file per role."),
    (0, "bullet",
     "Flag (don't auto-convert) any row filter that isn't a simple equality/"
     "IN-list pattern - same philosophy already used for calculated members "
     "today."),

    (1, "heading", "2. Perspectives"),
    (0, "para", "Feasible, but low practical value on Direct Lake today."),
    (0, "bullet",
     "AMO exposes Cube.Perspectives with per-object visibility flags (which "
     "measures/dimensions/hierarchies are included)."),
    (0, "bullet",
     "TMDL supports a perspective object listing included table/column/measure "
     "references - a straightforward mechanical 1:1 mapping."),
    (0, "bullet",
     "Caveat: current Fabric documentation indicates Direct Lake reports don't yet "
     "fully honor perspectives to restrict the visible field list. Still worth "
     "generating (future-proof, works when opened via XMLA/Tabular Editor), but "
     "should be flagged as \"generated but not yet enforced in the Fabric portal "
     "UI\" so expectations are set correctly."),

    (1, "heading", "3. Translations"),
    (0, "para", "Partial support."),
    (0, "bullet",
     "AMO exposes Database.Translations / Cube.Translations with translated "
     "captions per culture. TMDL supports a culture object with translation "
     "overrides - mechanical mapping is straightforward extraction-wise."),
    (0, "bullet",
     "Caveat: current Fabric documentation indicates Direct Lake models have "
     "limited multi-culture/translation support (English reliably works; other "
     "locales are inconsistent)."),
    (0, "bullet",
     "Recommendation: extract and generate the TMDL culture files (low cost, "
     "harmless if unsupported), but document the caveat rather than promising "
     "full support."),

    (1, "heading", "4. Actions"),
    (0, "para", "Not feasible as a direct port."),
    (0, "bullet",
     "SSAS \"Actions\" (drillthrough, reporting, standard MDX actions triggered "
     "from client tools like Excel/SSMS) have no equivalent concept in a Power "
     "BI/Fabric semantic model - Fabric's analog is a Power BI report's "
     "drillthrough page, which lives in the report (PBIX/PBIR), not the semantic "
     "model (TMDL) this tool generates."),
    (0, "bullet",
     "Since this tool only produces the semantic model, not reports, Actions "
     "can't be converted here at all. At best, we could extract and list them in "
     "the Migration Conversion Report as a manual note (\"recreate as "
     "report-level drillthrough pages\"), similar to how out-of-scope items are "
     "already flagged."),

    (1, "heading", "5. Custom Rollup Formulas / Unary Operators"),
    (0, "para", "Workaround only, high effort."),
    (0, "bullet",
     "These are MDX-level per-member override formulas on parent-child/ragged "
     "hierarchies (e.g. \"+/-\" unary operators for account rollups). DAX has no "
     "per-member override concept in the same way."),
    (0, "bullet",
     "The closest approximation is CALCULATE-based conditional aggregation logic "
     "keyed off a physical \"sign\"/\"rollup operator\" column materialized in the "
     "Lakehouse table - similar to the parent-child hierarchy path-precompute "
     "pattern already partially handled."),
    (0, "bullet",
     "This would need cube-specific DAX authored by a human, not something this "
     "tool can safely auto-generate. Realistic scope: extract and flag with the "
     "specific unary operator values per member, not full conversion."),

    (1, "heading", "6. Write-Back"),
    (0, "para", "Not supported, out of scope."),
    (0, "bullet",
     "Fabric Direct Lake and Import semantic models are read-only from Power "
     "BI's perspective (no WriteEnabled partition equivalent in Fabric's public "
     "model)."),
    (0, "bullet",
     "This one genuinely can't be added; best option is to keep flagging it as "
     "unsupported with no path forward until Fabric adds native write-back."),

    (1, "heading", "7. MDX SCOPE Assignments"),
    (0, "para", "Partial, case-by-case, highest effort."),
    (0, "bullet",
     "SCOPE(...) = ...; END SCOPE; blocks in an MDX script are essentially "
     "calculation overrides scoped to specific tuples/subcubes - there's no DAX "
     "equivalent construct at all (DAX has no concept of \"scope of the cube "
     "space\")."),
    (0, "bullet",
     "Some simple SCOPE patterns (e.g. currency conversion, a single measure "
     "override for a slice) can be manually re-expressed as DAX measures using "
     "CALCULATE/SWITCH with explicit filter contexts, but there's no general "
     "mechanical translation."),
    (0, "bullet",
     "Realistic scope: extract the raw MDX script text and calculation "
     "properties (CalculationProperty collection) and include them verbatim in "
     "MANUAL_TRANSLATION_REQUIRED.md for expert review - same flagging pattern "
     "as calculated members today, not auto-conversion."),

    (1, "heading", "Recommended Prioritization If Proceeding"),
    (0, "bullet",
     "1. Row-Level Security - highest value, most mechanical, extractor + TMDL "
     "generator work only."),
    (0, "bullet",
     "2. Perspectives - low effort, do alongside RLS, with a documented caveat "
     "about Direct Lake enforcement."),
    (0, "bullet",
     "3. Translations - low effort, same caveat pattern."),
    (0, "bullet",
     "4. Everything else (Actions, custom rollups, write-back, SCOPE) - stay in "
     "\"extract & flag for manual review\" territory rather than promising "
     "conversion, consistent with how calculated members/KPIs are already "
     "handled."),
]


def build_document():
    doc = Document()

    title = doc.add_heading(
        "Feasibility Review: Currently-Unsupported SSAS Features", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Row-Level Security, Actions, Perspectives, Translations, custom "
        "rollups, write-back, and MDX SCOPE assignments"
    )
    run.italic = True
    run.font.size = Pt(13)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Status: for review only - no implementation has been scoped or "
        "committed to yet. Revisit this document when ready to prioritize "
        "any of these items."
    )
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=len(SUMMARY_ROWS), cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row_values in enumerate(SUMMARY_ROWS):
        for c, value in enumerate(row_values):
            cell = table.cell(r, c)
            cell.text = value
            if r == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    doc.add_paragraph()

    doc.add_heading("Details", level=1)
    for level, kind, text in SECTIONS:
        if kind == "heading":
            doc.add_heading(text, level=level)
        elif kind == "para":
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.bold = True
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_document()

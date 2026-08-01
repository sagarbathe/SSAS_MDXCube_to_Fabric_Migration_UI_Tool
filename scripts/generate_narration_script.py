"""
One-off generator for docs/Video_Narration_Script.docx - a narration script
for a demo/overview video of the SSAS -> Fabric Migration UI Tool.

Not part of the pipeline; run manually whenever the script content below
needs to be regenerated after a real UI/feature change:
    .venv-ui\\Scripts\\python.exe scripts\\generate_narration_script.py
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(REPO_ROOT, "docs", "Video_Narration_Script.docx")

# Each entry: (heading_level, "action"|"narration"|"title", text)
SCRIPT = [
    (0, "title", "SSAS Multidimensional -> Microsoft Fabric Migration Tool"),
    (0, "subtitle", "Demo Video Narration Script"),

    # ------------------------------------------------------------------
    (1, "heading", "1. Purpose / Objective"),
    (0, "action", "[Screen: title slide / app open on the Read Me tab]"),
    (0, "narration",
     "Organizations running legacy SSAS Multidimensional cubes need a repeatable, "
     "low-risk path to Microsoft Fabric. Rebuilding a semantic model by hand is slow "
     "and error-prone - measures, hierarchies, relationships, and data types all have "
     "to be re-derived correctly from the cube."),
    (0, "narration",
     "This tool automates that end-to-end. It connects to a running SSAS cube, "
     "extracts its full metadata, analyzes it against Fabric's Direct Lake "
     "requirements, generates a ready-to-deploy Power BI semantic model plus starter "
     "Fabric notebook scripts, migrates the underlying star-schema data into a Fabric "
     "Lakehouse, and deploys the semantic model - all through a simple browser-based "
     "UI, no code required to run it."),
    (0, "narration",
     "It's generic: point it at any SSAS Multidimensional server and database, and it "
     "will extract and convert that cube's structure - it isn't hard-coded to one demo "
     "cube."),
    (0, "narration",
     "The tool is split into two phases that can run on two completely different "
     "machines: Phase 1 only needs to reach the on-prem SSAS server and SQL Server - "
     "no Fabric connectivity at all. Phase 2 only needs to reach Fabric - no SSAS "
     "connectivity at all. That split matters for organizations where the on-prem "
     "network and the Fabric-connected network aren't the same machine."),

    # ------------------------------------------------------------------
    (1, "heading", "2. Limitations Overview"),
    (0, "action", "[Screen: Read Me tab, scrolled to the Limitations excerpt]"),
    (0, "narration",
     "Before we walk through the steps, a few honest limitations to set expectations."),
    (0, "narration",
     "MDX calculated members, KPIs, and parent-child hierarchies are never "
     "auto-translated - MDX and DAX aren't mechanically equivalent, so instead of "
     "risking a silently wrong number, the tool extracts and lists them for a human "
     "to hand-author as DAX."),
    (0, "narration",
     "Semi-additive measures and many-to-many relationships are flagged for manual "
     "review rather than converted automatically. And Row-Level Security, Actions, "
     "Perspectives, Translations, custom rollups, write-back, and MDX SCOPE "
     "assignments aren't extracted at all today."),
    (0, "narration",
     "On the data side, the built-in 'migrate data' step needs one machine that can "
     "reach both the on-prem SQL Server and Fabric at the same time; for isolated "
     "networks there's an offline export-then-upload bridge instead, but for very "
     "large fact tables a Fabric-native pipeline or Spark notebook - which this tool "
     "also generates a starter script for - will scale better."),
    (0, "narration",
     "And two Fabric-specific quirks worth knowing up front: switching an "
     "already-deployed model between Direct Lake and Import mode isn't supported "
     "in-place by Fabric, so the tool automatically deletes and recreates the item "
     "when that happens - and an Import-mode model's very first refresh always needs "
     "a one-time manual credential binding in the Fabric portal before it succeeds. "
     "Both are called out clearly in the app when they happen, with exact remediation "
     "steps."),
    (0, "narration",
     "With that context set, let's walk through the app itself, one step at a time, "
     "in the order you'd actually click through them."),

    # ------------------------------------------------------------------
    (1, "heading", "3. Getting Started: Configuration"),
    (0, "action", "[Screen: click the 'Config' tab]"),
    (0, "narration",
     "Before running any step, we fill in connection details once on the Config tab: "
     "the on-prem SSAS server and cube name, the on-prem SQL Server and database, and "
     "the Fabric side - tenant ID, service principal client ID and secret, workspace "
     "ID, and the semantic model name. These get saved to a local .env file, exactly "
     "like the equivalent command-line tool would read with --env-file."),
    (0, "action", "[Screen: click 'Load from file' to show values populate; scroll to 'Run settings']"),
    (0, "narration",
     "Down in Run settings there's also the output folder for generated artifacts, "
     "and the Python executable used to actually run the pipeline steps - this "
     "defaults automatically to the right environment, so most users never need to "
     "touch it."),

    # ------------------------------------------------------------------
    (1, "heading", "4. Phase 1 Walkthrough: On-Prem (SSAS + SQL Server)"),
    (0, "action", "[Screen: click the 'Phase 1' tab]"),
    (0, "narration",
     "Now let's go through Phase 1. Every step here only needs the on-prem network - "
     "no Fabric connectivity required yet."),

    (2, "heading", "Step 1: Extract cube metadata"),
    (0, "action", "[Screen: click 'Run this step' under Step 1]"),
    (0, "narration",
     "First, Extract cube metadata. This connects to our SSAS Multidimensional cube "
     "using AMO - Windows Integrated Authentication only, so the account running this "
     "needs the Server Administrator role on the SSAS instance. It pulls the full "
     "metadata: dimensions, measures, hierarchies, calculated members, KPIs, and the "
     "data source view schema - and writes it all to cube_metadata.json."),
    (0, "action", "[Screen: watch the live log stream fill the scrollable progress box as it runs]"),
    (0, "narration",
     "Notice the progress box streaming live output as the step runs, wrapped so long "
     "lines are easy to read, with a spinner showing work is actively in progress."),

    (2, "heading", "Step 2: Analyze Direct Lake feasibility"),
    (0, "action", "[Screen: click 'Run this step' under Step 2]"),
    (0, "narration",
     "Next, Analyze Direct Lake feasibility. This is a pure offline analysis - no "
     "network calls - that checks the extracted metadata against Fabric's Direct Lake "
     "constraints and recommends Direct Lake or Import mode for this cube, with "
     "itemized reasons for any fallback, like parent-child hierarchies or semi-additive "
     "measures."),
    (0, "action", "[Screen: expand 'View / download feasibility_report.json']"),
    (0, "narration",
     "Right underneath, we can expand 'View / download feasibility report' to see the "
     "recommendation in a friendly table, and download it as an Excel workbook to "
     "share with the data modeling team - no need to read raw JSON."),

    (2, "heading", "Step 3: Generate TMDL + notebook scripts"),
    (0, "action", "[Screen: click 'Run this step' under Step 3]"),
    (0, "narration",
     "Step 3, Generate TMDL and notebook scripts. Also fully offline. This produces a "
     "TMDL semantic model folder - tables, columns, relationships, DAX measures "
     "translated from MDX, hierarchies - plus starter PySpark notebook scripts for "
     "teams that want a Spark-based load path. Anything that can't be safely "
     "auto-translated, like calculated members, is flagged instead of guessed at."),

    (2, "heading", "Step 4: Generate MIGRATION_REPORT.md"),
    (0, "action", "[Screen: click 'Run this step' under Step 4]"),
    (0, "narration",
     "Finally in Phase 1, Generate the Migration Report. This is the single "
     "plain-language document stating exactly what was converted automatically and "
     "what needs manual attention, with a suggested alternative for everything that "
     "wasn't converted."),
    (0, "action", "[Screen: expand 'View / download MIGRATION_REPORT.md', show the Word download button]"),
    (0, "narration",
     "Just like the feasibility report, we can view it right here or download it as a "
     "Word document to review and share before moving on to Phase 2. This is exactly "
     "the checkpoint to pause at - review this report before touching Fabric at all."),

    # ------------------------------------------------------------------
    (1, "heading", "5. Phase 2 Walkthrough: Fabric-Connected"),
    (0, "action", "[Screen: click the 'Phase 2' tab]"),
    (0, "narration",
     "Now Phase 2. From here on, every step only needs Fabric connectivity - no SSAS "
     "access required."),

    (2, "heading", "Step 5: Deploy / find Lakehouse"),
    (0, "action", "[Screen: choose 'Use an existing Lakehouse', click 'Refresh list of Lakehouses']"),
    (0, "narration",
     "Step 5, Deploy or find the Lakehouse. We can pick an existing Lakehouse from a "
     "live dropdown pulled straight from the Fabric workspace, or create a brand new "
     "one by name."),
    (0, "action", "[Screen: click 'Deploy / find Lakehouse']"),
    (0, "narration",
     "Clicking Deploy or find Lakehouse creates it if it doesn't exist yet, or reuses "
     "it if it does, and patches the generated TMDL with its real SQL analytics "
     "endpoint so the semantic model knows exactly where to find its data."),

    (2, "heading", "Step 6: Migrate data into the Lakehouse"),
    (0, "action", "[Screen: show the 'direct' vs 'offline' radio choice]"),
    (0, "narration",
     "Step 6, migrate the star-schema data. There are two paths depending on network "
     "reachability. Direct migration is the simplest: this one machine reaches both "
     "SQL Server and Fabric, and clicking 'Migrate data now' extracts the tables and "
     "writes them straight into the Lakehouse as Delta tables via OneLake."),
    (0, "action", "[Screen: switch the radio to 'offline transfer', show the two-step flow]"),
    (0, "narration",
     "For isolated or air-gapped networks, there's an offline transfer option instead: "
     "export the tables to local Delta files on the on-prem side first, manually "
     "transfer that folder over using whatever secure method your organization "
     "already approves, and then upload it to the Lakehouse from the Fabric-connected "
     "machine - no direct network path between the two is ever required."),
    (0, "action", "[Screen: expand 'Other ways to migrate data into Fabric']"),
    (0, "narration",
     "For larger, production-scale data volumes, there's also a reference list of "
     "Fabric-native alternatives - Mirroring, Open Mirroring, Data Factory, Dataflows "
     "Gen2, and Spark notebooks - worth reviewing if this simple built-in path won't "
     "scale for your fact table sizes."),

    (2, "heading", "Step 7: Deploy semantic model"),
    (0, "action", "[Screen: click 'Deploy semantic model']"),
    (0, "narration",
     "And finally, Step 7, deploy the semantic model. This creates or updates the "
     "Semantic Model item in the target Fabric workspace from the generated TMDL. If "
     "a previously deployed model's storage mode needs to flip between Direct Lake "
     "and Import, Fabric doesn't support that in place, so the tool automatically "
     "deletes and recreates the item instead of leaving it broken."),
    (0, "narration",
     "After deploying, it automatically triggers a refresh and waits for it to "
     "finish, so the model is immediately ready for report-building - without this, "
     "Fabric would show stale or missing data until someone manually clicked 'Refresh "
     "now' in the portal. On an Import-mode model's very first deploy, this refresh "
     "is expected to warn rather than fail outright, until credentials are bound once "
     "in the Fabric portal - that's called out clearly in the log output."),

    # ------------------------------------------------------------------
    (1, "heading", "6. Reports Tab Overview"),
    (0, "action", "[Screen: click the 'Reports' tab]"),
    (0, "narration",
     "One last stop - the Reports tab. Think of this as a single, permanent home for "
     "every document this tool generates, so nobody has to go hunting through the "
     "output folder on disk to find them. It always reflects whatever the most recent "
     "run produced, regardless of which Phase 1 step you happened to trigger it from."),

    (2, "heading", "Migration Conversion Report"),
    (0, "action", "[Screen: show the Migration Conversion Report rendered at the top of the tab]"),
    (0, "narration",
     "At the top is the Migration Conversion Report, rendered right on the page in "
     "plain language - the same one we generated back in Phase 1 Step 4 - showing "
     "exactly what converted automatically and what needs manual attention."),
    (0, "action", "[Screen: click the 'Download as Word' button]"),
    (0, "narration",
     "There's a one-click download as a formatted Word document here too, so it's "
     "easy to attach to an email or share with a reviewer who doesn't have access to "
     "this tool."),

    (2, "heading", "Feasibility Report"),
    (0, "action", "[Screen: expand the 'feasibility_report.json' section, show the friendly table]"),
    (0, "narration",
     "Next, an expandable section for the feasibility report - the Direct Lake versus "
     "Import recommendation and reasons from Phase 1 Step 2 - shown as a readable "
     "table rather than raw JSON, with the same Excel download option we saw earlier."),

    (2, "heading", "Manual Translation Required"),
    (0, "action", "[Screen: expand the 'MANUAL_TRANSLATION_REQUIRED.md' section]"),
    (0, "narration",
     "And finally, an expandable section for Manual Translation Required - the list "
     "of MDX calculated members, KPIs, and other constructs that couldn't be safely "
     "auto-converted, each with its original MDX text included, ready to hand off for "
     "DAX authoring. This one only appears if the cube actually has anything that "
     "needed flagging - simpler cubes with no calculated members won't produce it at "
     "all. It's downloadable as Word too, just like the other two."),
    (0, "narration",
     "So in short: three reports, one tab, always up to date, and every one of them "
     "exportable in a business-friendly format - no digging through JSON or Markdown "
     "files on disk required."),

    (0, "action", "[Screen: return to title slide]"),
    (0, "narration",
     "That's the full journey - from a live SSAS cube, through feasibility analysis "
     "and semantic model generation, to a deployed, refreshed Direct Lake or Import "
     "model in Microsoft Fabric, all driven from one browser tab."),
]


def build_document():
    doc = Document()

    for level, kind, text in SCRIPT:
        if kind == "title":
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(16)
        elif kind == "heading":
            doc.add_heading(text, level=level)
        elif kind == "action":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif kind == "narration":
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(8)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_document()

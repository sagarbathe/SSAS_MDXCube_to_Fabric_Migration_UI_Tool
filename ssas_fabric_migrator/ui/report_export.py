"""
Converts the pipeline's generated Markdown reports and feasibility JSON into
Word (.docx) and Excel (.xlsx) files for users who'd rather review/share
findings in Office than raw Markdown/JSON. Used only by the Streamlit UI
(app.py) - not part of the CLI/orchestrator pipeline itself.

Kept deliberately simple: this is a "good enough for review" converter, not
a full Markdown/CommonMark implementation. It handles exactly what this
tool's own generated reports use: #/##/### headings, GFM pipe tables,
"- " bullets (including "- [ ] " checklist items), and plain paragraphs
with **bold** spans.
"""
from __future__ import annotations

import io
import re

SEVERITY_FILL = {
    "blocking": "FFC7CE",  # light red
    "warning": "FFEB9C",  # light amber
    "info": "C6E0B4",  # light green
}


def _add_markdown_runs(paragraph, text: str):
    """Splits **bold** spans out of a line of text into separate runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _parse_pipe_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


_SEPARATOR_RE = re.compile(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?$")


def markdown_to_docx_bytes(markdown_text: str, doc_title: str) -> bytes:
    """Renders this tool's generated Markdown reports as a Word document.
    Returns the .docx file content as bytes, ready for st.download_button."""
    from docx import Document

    document = Document()
    document.add_heading(doc_title, level=0)

    lines = markdown_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        # GFM pipe table: a header row, a separator row, then 0+ data rows.
        if stripped.startswith("|") and i + 1 < n and _SEPARATOR_RE.match(lines[i + 1].strip()):
            header_cells = _parse_pipe_row(stripped)
            j = i + 2
            data_rows = []
            while j < n and lines[j].strip().startswith("|"):
                data_rows.append(_parse_pipe_row(lines[j]))
                j += 1
            table = document.add_table(rows=1, cols=len(header_cells))
            table.style = "Light Grid Accent 1"
            for col_idx, cell_text in enumerate(header_cells):
                table.rows[0].cells[col_idx].paragraphs[0].add_run(cell_text).bold = True
            severity_col = None
            for col_idx, cell_text in enumerate(header_cells):
                if cell_text.strip().lower() == "severity":
                    severity_col = col_idx
            for row_cells in data_rows:
                row = table.add_row()
                for col_idx, cell_text in enumerate(row_cells):
                    if col_idx < len(row.cells):
                        row.cells[col_idx].text = cell_text
                if severity_col is not None and severity_col < len(row_cells):
                    fill = SEVERITY_FILL.get(row_cells[severity_col].strip().lower())
                    if fill:
                        _shade_cell(row.cells[severity_col], fill)
            i = j
            continue

        checklist_match = re.match(r"^-\s*\[( |x|X)\]\s*(.*)$", stripped)
        if checklist_match:
            box = "\u2611" if checklist_match.group(1).lower() == "x" else "\u2610"
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, f"{box} {checklist_match.group(2)}")
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, bullet_match.group(1))
            i += 1
            continue

        p = document.add_paragraph()
        _add_markdown_runs(p, stripped)
        i += 1

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _shade_cell(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def feasibility_json_to_excel_bytes(feasibility_report: dict) -> bytes:
    """Renders feasibility_report.json (cubes -> recommended mode + findings)
    as a two-sheet Excel workbook: a per-cube Summary sheet and a flattened,
    severity-highlighted Findings sheet. Returns .xlsx bytes for
    st.download_button."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    summary_ws = wb.active
    summary_ws.title = "Summary"
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4472C4")

    summary_headers = ["Cube", "Recommended Mode", "Total Findings", "Blocking Findings"]
    summary_ws.append(summary_headers)
    for col in range(1, len(summary_headers) + 1):
        c = summary_ws.cell(row=1, column=col)
        c.font = header_font
        c.fill = header_fill

    findings_ws = wb.create_sheet("Findings")
    findings_headers = ["Cube", "Severity", "Scope", "Message"]
    findings_ws.append(findings_headers)
    for col in range(1, len(findings_headers) + 1):
        c = findings_ws.cell(row=1, column=col)
        c.font = header_font
        c.fill = header_fill

    for cube in feasibility_report.get("cubes", []):
        name = cube.get("cube_name", "")
        mode = cube.get("recommended_mode", "")
        findings = cube.get("findings", [])
        blocking_count = sum(1 for f in findings if f.get("severity", "").lower() == "blocking")
        summary_ws.append([name, mode, len(findings), blocking_count])

        for finding in findings:
            severity = finding.get("severity", "")
            row = [name, severity, finding.get("scope", ""), finding.get("message", "")]
            findings_ws.append(row)
            fill_color = SEVERITY_FILL.get(severity.lower())
            if fill_color:
                row_idx = findings_ws.max_row
                fill = PatternFill("solid", fgColor=fill_color)
                for col in range(1, len(findings_headers) + 1):
                    findings_ws.cell(row=row_idx, column=col).fill = fill

    for ws, widths in (
        (summary_ws, [30, 18, 15, 18]),
        (findings_ws, [30, 12, 40, 80]),
    ):
        for idx, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = width
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
